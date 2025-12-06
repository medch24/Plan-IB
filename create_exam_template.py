#!/usr/bin/env python3
"""
Script pour créer un template d'examen Word avec les bonnes spécifications :
- Marges : 1.5 cm (gauche, droite, haut, bas)
- Interligne : 1.5
- Format professionnel
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Check for tag existence, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # List over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # Check for tag existence, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # Add attributes for the edge
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_exam_template():
    """Créer un template d'examen avec les spécifications demandées"""
    
    doc = Document()
    
    # === CONFIGURATION DES MARGES : 1.5 CM ===
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # === EN-TÊTE DE L'EXAMEN ===
    # Tableau pour l'en-tête
    table_header = doc.add_table(rows=2, cols=3)
    table_header.style = 'Table Grid'
    
    # Ligne 1
    cells_row1 = table_header.rows[0].cells
    
    # Cellule 1,1 - Titre de l'examen
    cell_11 = cells_row1[0]
    p = cell_11.paragraphs[0]
    p.add_run('Examen').bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Cellule 1,2 - Matière (placeholder)
    cell_12 = cells_row1[1]
    p = cell_12.paragraphs[0]
    run = p.add_run('{Matiere}')
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Cellule 1,3 - Classe
    cell_13 = cells_row1[2]
    p = cell_13.paragraphs[0]
    run = p.add_run('Classe : {Classe}')
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ligne 2
    cells_row2 = table_header.rows[1].cells
    
    # Cellule 2,1 - Durée
    cell_21 = cells_row2[0]
    p = cell_21.paragraphs[0]
    p.add_run('Durée: {Duree}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Cellule 2,2 - Enseignant
    cell_22 = cells_row2[1]
    p = cell_22.paragraphs[0]
    p.add_run('Enseignant: {Enseignant}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Cellule 2,3 - Semestre
    cell_23 = cells_row2[2]
    p = cell_23.paragraphs[0]
    p.add_run('Semestre: {Semestre}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date en-dessous du tableau
    doc.add_paragraph()
    p_date = doc.add_paragraph('Date : {Date}')
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # === INFORMATIONS ÉLÈVE ===
    doc.add_paragraph()
    p_nom = doc.add_paragraph('Nom et prénom : ............................................................')
    p_nom.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # === NOTATION ===
    doc.add_paragraph()
    table_note = doc.add_table(rows=1, cols=2)
    table_note.style = 'Table Grid'
    
    cells_note = table_note.rows[0].cells
    cells_note[0].text = 'Note : .......... / 30'
    cells_note[1].text = 'Observations : .......................................................'
    
    # === ESPACE POUR EXERCICES ===
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Titre des exercices
    p_exercices = doc.add_paragraph('EXERCICES')
    p_exercices.runs[0].bold = True
    p_exercices.runs[0].font.size = Pt(12)
    p_exercices.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_exercices.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    doc.add_paragraph()
    
    # === PLACEHOLDER POUR LE CONTENU DES EXERCICES ===
    # Ce sera rempli dynamiquement par docxtemplater
    p_content = doc.add_paragraph('{Exercices}')
    p_content.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # === CONFIGURATION GLOBALE DE L'INTERLIGNE ===
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # === SAUVEGARDE ===
    output_path = '/home/user/webapp/public/Template_Examen_Ministere_New.docx'
    doc.save(output_path)
    print(f"✅ Template créé avec succès : {output_path}")
    print("📐 Marges : 1.5 cm (tous côtés)")
    print("📏 Interligne : 1.5")
    print("✨ Formatage : Professionnel")

if __name__ == '__main__':
    create_exam_template()
